#!/usr/bin/env python3
"""
Generate a professionally formatted Word (.docx) document from a Markdown transcript.

Supports embedded images (![alt](path)) and timestamp annotations in headings.

Usage:
    python generate_docx.py --input transcript.md --output transcript.docx \
        --title "Document Title" --author "Speaker Name" --source "https://..."
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    shading.append(shading_elem)


def create_docx(
    md_content: str,
    output_path: str,
    title: str = "Transcript",
    author: str = "",
    source: str = "",
    date: str = "",
    language: str = "",
    duration: str = "",
    base_dir: str | Path | None = None,
):
    """Create a formatted Word document from markdown content.

    Args:
        base_dir: Base directory for resolving relative image paths.
                  Defaults to the parent directory of output_path.
    """
    if base_dir is None:
        base_dir = Path(output_path).parent
    base_dir = Path(base_dir)

    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(12)
    font.name = "Calibri"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(3)

    # Title style
    title_style = doc.styles["Title"]
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.name = "Calibri"
    title_style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Heading styles
    for level, size in [(1, 18), (2, 15), (3, 13)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.size = Pt(size)
        h_style.font.bold = True
        h_style.font.name = "Calibri"
        h_style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        h_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        h_style.paragraph_format.space_before = Pt(18)
        h_style.paragraph_format.space_after = Pt(8)

    # --- Title Page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = "Calibri"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Subtitle / metadata line
    if author or date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(24)
        meta_parts = []
        if author:
            meta_parts.append(f"Speaker: {author}")
        if date:
            meta_parts.append(f"Date: {date}")
        if language:
            meta_parts.append(f"Language: {language}")
        if duration:
            meta_parts.append(f"Duration: {duration}")
        run = p.add_run("  |  ".join(meta_parts))
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    if source:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(8)
        run = p.add_run(f"Source: {source}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    # Page break after title page
    doc.add_page_break()

    # --- Parse and render Markdown ---
    lines = md_content.split("\n")
    i = 0

    # Skip YAML frontmatter if present
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        i += 1  # Skip closing ---

    # Skip metadata block (> **Source**: ...)
    while i < len(lines) and (lines[i].strip().startswith(">") or lines[i].strip() == ""):
        i += 1

    # Skip horizontal rules at the top
    while i < len(lines) and lines[i].strip() in ["---", "***", "___", ""]:
        i += 1

    while i < len(lines):
        line = lines[i].strip()

        # Skip horizontal rules
        if line in ["---", "***", "___"]:
            i += 1
            continue

        # Skip metadata lines starting with >
        if line.startswith("> **"):
            i += 1
            continue

        # Headings
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            heading_text = line.lstrip("#").strip()
            if level <= 3 and heading_text:
                # Skip "Table of Contents" heading - we'll generate our own
                if heading_text.lower() in ["table of contents", "目录"]:
                    i += 1
                    # Skip TOC content
                    while i < len(lines) and (lines[i].strip().startswith("-") or lines[i].strip().startswith("*") or lines[i].strip() == ""):
                        if lines[i].strip() == "" and i + 1 < len(lines) and lines[i + 1].strip().startswith("#"):
                            break
                        i += 1
                    continue

                # Extract timestamp suffix like `[00:01:23]` from heading
                ts_match = re.search(r"\s*`\[(\d{2}:\d{2}:\d{2})\]`\s*$", heading_text)
                timestamp_str = ""
                if ts_match:
                    timestamp_str = ts_match.group(1)
                    heading_text = heading_text[:ts_match.start()].strip()

                h = doc.add_heading(heading_text, level=min(level, 3))

                # Append timestamp as styled run
                if timestamp_str:
                    run = h.add_run(f"  [{timestamp_str}]")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
                    run.font.bold = False
            i += 1
            continue

        # Images: ![alt text](path)
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if img_match:
            alt_text = img_match.group(1)
            img_src = img_match.group(2)
            # Resolve relative path against base_dir
            img_path = base_dir / img_src
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(5.5))
                    # Add caption
                    if alt_text:
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cap.add_run(alt_text)
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
                        run.font.italic = True
                except Exception as e:
                    # If image embedding fails, add as text
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Image: {alt_text}]")
                    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
                    run.font.italic = True
            i += 1
            continue

        # Blockquotes (as indented italic text)
        if line.startswith(">"):
            quote_text = line.lstrip(">").strip()
            if quote_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.5)
                run = p.add_run(quote_text)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # List items
        if re.match(r"^[-*]\s", line) or re.match(r"^\d+\.\s", line):
            if re.match(r"^[-*]\s", line):
                text = re.sub(r"^[-*]\s+", "", line)
                p = doc.add_paragraph(style="List Bullet")
            else:
                text = re.sub(r"^\d+\.\s+", "", line)
                p = doc.add_paragraph(style="List Number")

            # Handle inline formatting
            _add_formatted_text(p, text)
            i += 1
            continue

        # Empty lines
        if not line:
            i += 1
            continue

        # Regular paragraphs - collect consecutive non-empty lines
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and not lines[i].strip().startswith(">") and lines[i].strip() not in ["---", "***", "___"]:
            para_lines.append(lines[i].strip())
            i += 1

        if para_lines:
            text = " ".join(para_lines)
            p = doc.add_paragraph()
            _add_formatted_text(p, text)
            continue

        i += 1

    # --- Footer ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(36)
    run = p.add_run("—  End of Transcript  —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
    run.font.italic = True

    if source:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Extracted from: {source}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)

    # --- Save ---
    doc.save(output_path)
    print(f"Document saved: {output_path}")


def _add_formatted_text(paragraph, text: str):
    """Add text with inline markdown formatting (bold, italic, code)."""
    # Pattern to match **bold**, *italic*, `code`, and plain text
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)"
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)


def main():
    parser = argparse.ArgumentParser(description="Generate formatted Word document from Markdown")
    parser.add_argument("--input", required=True, help="Input Markdown file")
    parser.add_argument("--output", required=True, help="Output .docx file")
    parser.add_argument("--title", default="Transcript", help="Document title")
    parser.add_argument("--author", default="", help="Speaker/author name")
    parser.add_argument("--source", default="", help="Source URL")
    parser.add_argument("--date", default="", help="Video date")
    parser.add_argument("--language", default="", help="Source language")
    parser.add_argument("--duration", default="", help="Video duration")
    parser.add_argument("--base-dir", default=None, help="Base directory for resolving relative image paths")
    args = parser.parse_args()

    md_path = Path(args.input)
    if not md_path.exists():
        print(f"Error: Input file not found: {md_path}", file=sys.stderr)
        sys.exit(1)

    md_content = md_path.read_text(encoding="utf-8")

    create_docx(
        md_content=md_content,
        output_path=args.output,
        title=args.title,
        author=args.author,
        source=args.source,
        date=args.date,
        language=args.language,
        duration=args.duration,
        base_dir=args.base_dir,
    )


if __name__ == "__main__":
    main()
